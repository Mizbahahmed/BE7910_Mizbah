# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""
import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt  # Should work after reinstalling
import seaborn as sns
from sklearn.ensemble import RandomForestClassifier
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score
from docx import Document  # Import the python-docx library for Word document generation

# Set paths
INPUT_DIR = "D:\\LSU_Spring 2025\\BE 7901\\Assignment 4\\DATA\\archive"
OUTPUT_DIR = "D:\\LSU_Spring 2025\\BE 7901\\Assignment 4\\Outputs"
ACCURACY_REPORT_PATH = os.path.join(OUTPUT_DIR, "Accuracy.docx")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Load data
def load_data():
    train = pd.read_csv(os.path.join(INPUT_DIR, "train.csv"))
    test = pd.read_csv(os.path.join(INPUT_DIR, "test.csv"))
    return train, test

# Feature engineering
def preprocess_data(df):
    df = df.copy()
    # Family size feature
    df['FamilySize'] = df['SibSp'] + df['Parch'] + 1
    # Title extraction
    df['Title'] = df['Name'].str.extract(' ([A-Za-z]+)\.', expand=False)
    df['Title'] = df['Title'].where(df['Title'].isin(['Mr', 'Miss', 'Mrs', 'Master']), 'Rare')
    # Drop columns
    return df.drop(['Name', 'Ticket', 'Cabin', 'SibSp', 'Parch'], axis=1)

# Main pipeline
def main():
    # Load and preprocess
    train, test = load_data()
    train = preprocess_data(train)
    test = preprocess_data(test)
    
    # Imputation
    imputer = SimpleImputer(strategy='median')
    for col in ['Age', 'Fare']:
        train[col] = imputer.fit_transform(train[[col]])
        test[col] = imputer.transform(test[[col]])
    
    # Encoding
    train['Sex'] = train['Sex'].map({'male': 0, 'female': 1})
    test['Sex'] = test['Sex'].map({'male': 0, 'female': 1})
    train = pd.get_dummies(train, columns=['Embarked', 'Title'])
    test = pd.get_dummies(test, columns=['Embarked', 'Title'])
    
    # Align features
    train_features = train.drop(['PassengerId', 'Survived'], axis=1)
    test_features = test.drop(['PassengerId'], axis=1)
    
    # Ensure matching columns
    missing_cols = set(train_features.columns) - set(test_features.columns)
    for col in missing_cols:
        test_features[col] = 0
    
    # Train-test split
    X_train, X_val, y_train, y_val = train_test_split(
        train_features, train['Survived'], test_size=0.2, random_state=42
    )
    
    # Model training
    scaler = StandardScaler()
    model = RandomForestClassifier(n_estimators=200, max_depth=7, random_state=42)
    model.fit(scaler.fit_transform(X_train), y_train)
    
    # Validation
    val_pred = model.predict(scaler.transform(X_val))
    val_accuracy = accuracy_score(y_val, val_pred)
    print(f"Validation Accuracy: {val_accuracy:.4f}")
    
    # Generate predictions
    test_pred = model.predict(scaler.transform(test_features))
    submission = pd.DataFrame({
        'PassengerId': test['PassengerId'],
        'Survived': test_pred
    })
    submission.to_csv(os.path.join(OUTPUT_DIR, 'submission.csv'), index=False)
    print("Submission file created successfully")

    # Write accuracy assessment to Word document
    write_accuracy_report(val_accuracy)

# Compare accuracy with ground truth labels
def compare_accuracy():
    # Load the submission and ground truth files
    submission = pd.read_csv(os.path.join(OUTPUT_DIR, 'submission.csv'))
    ground_truth = pd.read_csv(os.path.join(OUTPUT_DIR, 'test_labels.csv'))  # Change to actual ground truth file path

    # Merge on 'PassengerId' to align predictions with true labels
    result = pd.merge(submission, ground_truth, on='PassengerId')

    # Calculate accuracy
    accuracy = accuracy_score(result['Survived_x'], result['Survived_y'])
    print(f"Accuracy: {accuracy:.4f}")
    return accuracy

# Function to write the accuracy report to a Word document
def write_accuracy_report(accuracy):
    doc = Document()
    doc.add_heading('Accuracy Assessment Report', 0)
    
    # Add accuracy to the report
    doc.add_paragraph(f"Validation Accuracy: {accuracy:.4f}")
    
    # Save the report to a Word document
    doc.save(ACCURACY_REPORT_PATH)
    print(f"Accuracy report saved at: {ACCURACY_REPORT_PATH}")

# Run the main pipeline and accuracy comparison
if __name__ == "__main__":
    main()  # Train model and generate submission
    compare_accuracy()  # Compare accuracy with ground truth labels
